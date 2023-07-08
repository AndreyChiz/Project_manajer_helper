import os

import openpyxl
from docx import Document


def parse_xlsx(self) -> list(list()):
    """Парсит таблицу xlsx

    Returns:
        list(): сырой список без учёта 'Разделов' и 'Подразделов'
    """
    wb = openpyxl.load_workbook(self.filename)
    ws = wb.active
    data = []
    for row in ws.iter_rows():
        row_data = [cell.value for cell in row if cell.value]
        if row_data:
            data


def build_list_new_text(raw_lst: list(list)) -> list(list()):
    """Собирает список замен.

    Returns:
        list(): список новых текстовых элементов для замены существующих текстовых элементов шаблона
    """
    for str_of_table in raw_lst:
        tmp_lst = []
        res_lst = []
        part = ''
        subpart = ''

        def what_is_str_of_table(x: list()) -> bool:
            return len(
                str_of_table) == 1 and str_of_table[0].startswith(x)

        is_part = what_is_str_of_table('Раздел')
        is_subpart = what_is_str_of_table('Подраздел')

        if is_part:
            part = ''
            subpart = ''
            part = str_of_table[0]
        elif is_subpart:
            subpart = ''
            subpart = str_of_table[0]
        else:
            tmp_lst.extend([part, subpart, *str_of_table])
        if tmp_lst:
            res_lst.append(tmp_lst)

    return res_lst


def build_list_old_text(path: str) -> list(list()):
    '''
    Возвращает список доступных для изменения параграфов из шаблона
    :return:
    '''
    doc = Document(path)
    paragraphs = [p.text for p in doc.paragraphs if p.text]
    return paragraphs


def doc_text_changer(input_file_path: str, output_file_puth: str, replacements: {str: str}) -> None:
    '''
    Ищет строки соответствующие ключу словаря 'replacements' и заменяет их на значение по этому ключу.
    :param input_file: путь шаблона
    :param output_file: путь для сохранения
    :param replacements: словарь {<Существующая строка>str: <На что заменить>:str}
    :return:
    '''
    document = Document(input_file_path)

    for paragraph in document.paragraphs:
        for old_text, new_text in replacements.items():
            if old_text in paragraph.text:
                tmp_run_text_list = []
                for run in paragraph.runs:
                    tmp_run_text_list.append(run.text)
                tmp_run_concatenation_string = ''.join(tmp_run_text_list)
                count = 0
                for run in paragraph.runs:
                    if count == 0:
                        run.text = tmp_run_concatenation_string.replace(old_text, new_text)
                        new_text = ''
                        count = 1
                    else:
                        run.text = ''
    document.save(output_file_puth)


# --------------------------------------------------
def create_output_file_puth(prefix: str = '', postfix: str = '', file_name_words_separator: str = '-' *,
                            name_parts_list: list() = [],
                            name_parts_order: list() = [1, 3, 2, 4]) -> str:
    '''

    :param prefix: перфикс имени файла
    :param postfix: постфикс
    :param name_parts_list: один список из списка списков new_text_list
    :param name_parts_order: список из индексов елементов из new_text_list, для добавления на позицию в имени файла
    соответствующую индексу списка ([1, 3, 2, 4]))
    :return: строка с путем сохранения и именем файла.
    Имя файла генерируется из префикса, нескольких строк из new_text_list, постфикса и засширения
    '''

    root_puth = os.getcwd()
    extension = '.docx'
    if len(name_parts_list) >= len(name_parts_order):
        generated_name_part = file_name_words_separator.join(
            [prefix, *[name_parts_list[i] for i in name_parts_order], postfix, ]) + extension
    else:
        print('Неверная очередь')
    puth = '\\'.join([root_puth, generated_name_part])
    return puth
